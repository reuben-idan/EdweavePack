from typing import Dict, Any
import json
from datetime import datetime
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
import io
import os

class ExportService:
    def __init__(self):
        self.styles = getSampleStyleSheet()
    
    def export_lesson_plan_pdf(self, curriculum_data: Dict[str, Any], format_type: str = "detailed") -> bytes:
        """Export curriculum as PDF lesson plan"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1  # Center
        )
        
        title = curriculum_data.get('curriculum_overview', 'Curriculum Plan')
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 12))
        
        # Overview
        story.append(Paragraph("Curriculum Overview", self.styles['Heading2']))
        overview = curriculum_data.get('curriculum_overview', 'No overview available')
        story.append(Paragraph(overview, self.styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Learning Objectives
        story.append(Paragraph("Learning Objectives", self.styles['Heading2']))
        objectives = curriculum_data.get('learning_objectives', [])
        for i, obj in enumerate(objectives, 1):
            story.append(Paragraph(f"{i}. {obj}", self.styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Weekly Modules
        for week in curriculum_data.get('weekly_modules', []):
            story.append(Paragraph(f"Week {week['week_number']}: {week['title']}", self.styles['Heading2']))
            story.append(Paragraph(f"Focus: {week.get('bloom_focus', 'N/A')}", self.styles['Normal']))
            
            # Content blocks
            for block in week.get('content_blocks', []):
                story.append(Paragraph(f"• {block['title']} ({block.get('estimated_duration', 0)} min)", self.styles['Normal']))
                story.append(Paragraph(f"  {block['description']}", self.styles['Normal']))
            
            story.append(Spacer(1, 12))
        
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    
    def export_lesson_plan_docx(self, curriculum_data: Dict[str, Any]) -> bytes:
        """Export curriculum as DOCX lesson plan"""
        doc = Document()
        
        # Title
        title = doc.add_heading(curriculum_data.get('curriculum_overview', 'Curriculum Plan'), 0)
        title.alignment = 1  # Center
        
        # Overview
        doc.add_heading('Curriculum Overview', level=1)
        overview = curriculum_data.get('curriculum_overview', 'No overview available')
        doc.add_paragraph(overview)
        
        # Learning Objectives
        doc.add_heading('Learning Objectives', level=1)
        objectives = curriculum_data.get('learning_objectives', [])
        for i, obj in enumerate(objectives, 1):
            doc.add_paragraph(f"{i}. {obj}", style='List Number')
        
        # Weekly Modules
        doc.add_heading('Weekly Schedule', level=1)
        for week in curriculum_data.get('weekly_modules', []):
            doc.add_heading(f"Week {week['week_number']}: {week['title']}", level=2)
            doc.add_paragraph(f"Bloom's Focus: {week.get('bloom_focus', 'N/A')}")
            
            # Learning outcomes
            if week.get('learning_outcomes'):
                doc.add_paragraph('Learning Outcomes:', style='Heading 3')
                for outcome in week['learning_outcomes']:
                    doc.add_paragraph(outcome, style='List Bullet')
            
            # Content blocks
            doc.add_paragraph('Content Blocks:', style='Heading 3')
            for block in week.get('content_blocks', []):
                p = doc.add_paragraph()
                p.add_run(f"{block['title']}").bold = True
                p.add_run(f" ({block.get('estimated_duration', 0)} minutes)")
                doc.add_paragraph(block['description'])
                
                if block.get('activities'):
                    doc.add_paragraph('Activities: ' + ', '.join(block['activities']))
            
            # Assessments
            if week.get('formative_assessments'):
                doc.add_paragraph('Formative Assessments:', style='Heading 3')
                for assessment in week['formative_assessments']:
                    doc.add_paragraph(assessment, style='List Bullet')
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    def generate_shareable_link(self, curriculum_id: int, user_id: int) -> str:
        """Generate shareable link for curriculum"""
        # In production, this would use a proper URL shortener or token system
        import hashlib
        import time
        
        # Create a simple hash-based link
        data = f"{curriculum_id}-{user_id}-{int(time.time())}"
        link_hash = hashlib.md5(data.encode()).hexdigest()[:12]
        
        # Return shareable URL (would be actual domain in production)
        return f"https://edweavepack.com/shared/{link_hash}"
    
    def export_project_rubric(self, project_data: Dict[str, Any]) -> bytes:
        """Export project rubric as PDF"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        story = []
        
        # Title
        story.append(Paragraph(f"Project Rubric: {project_data['title']}", self.styles['Heading1']))
        story.append(Spacer(1, 12))
        
        # Description
        story.append(Paragraph("Project Description", self.styles['Heading2']))
        story.append(Paragraph(project_data['description'], self.styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Scaffolds
        story.append(Paragraph("Project Steps", self.styles['Heading2']))
        for scaffold in project_data.get('scaffolds', []):
            story.append(Paragraph(f"Step {scaffold['step']}: {scaffold['task']}", self.styles['Heading3']))
            story.append(Paragraph(f"Support: {scaffold['support']}", self.styles['Normal']))
            story.append(Spacer(1, 6))
        
        # Rubric table
        story.append(Paragraph("Grading Rubric", self.styles['Heading2']))
        
        rubric = project_data.get('rubric', {})
        if rubric:
            # Create table data
            table_data = [['Criteria', 'Excellent', 'Good', 'Needs Improvement']]
            
            for criteria, levels in rubric.items():
                row = [criteria.replace('_', ' ').title()]
                row.append(levels.get('excellent', ''))
                row.append(levels.get('good', ''))
                row.append(levels.get('needs_improvement', levels.get('needs_work', '')))
                table_data.append(row)
            
            # Create and style table
            table = Table(table_data, colWidths=[1.5*Inches, 1.5*Inches, 1.5*Inches, 1.5*Inches])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(table)
        
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    
    def export_assessment_report(self, assessment_data: Dict[str, Any], student_results: list) -> bytes:
        """Export assessment results as PDF report"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        story = []
        
        # Title
        story.append(Paragraph("Assessment Report", self.styles['Heading1']))
        story.append(Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}", self.styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Assessment overview
        story.append(Paragraph("Assessment Overview", self.styles['Heading2']))
        story.append(Paragraph(f"Title: {assessment_data.get('title', 'N/A')}", self.styles['Normal']))
        story.append(Paragraph(f"Total Points: {assessment_data.get('total_points', 'N/A')}", self.styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Results summary
        if student_results:
            story.append(Paragraph("Results Summary", self.styles['Heading2']))
            
            total_students = len(student_results)
            avg_score = sum(r.get('percentage', 0) for r in student_results) / total_students
            
            story.append(Paragraph(f"Total Students: {total_students}", self.styles['Normal']))
            story.append(Paragraph(f"Average Score: {avg_score:.1f}%", self.styles['Normal']))
            
            # Performance distribution
            excellent = len([r for r in student_results if r.get('percentage', 0) >= 90])
            good = len([r for r in student_results if 80 <= r.get('percentage', 0) < 90])
            satisfactory = len([r for r in student_results if 70 <= r.get('percentage', 0) < 80])
            needs_improvement = len([r for r in student_results if r.get('percentage', 0) < 70])
            
            story.append(Paragraph(f"Excellent (90%+): {excellent} students", self.styles['Normal']))
            story.append(Paragraph(f"Good (80-89%): {good} students", self.styles['Normal']))
            story.append(Paragraph(f"Satisfactory (70-79%): {satisfactory} students", self.styles['Normal']))
            story.append(Paragraph(f"Needs Improvement (<70%): {needs_improvement} students", self.styles['Normal']))
        
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()